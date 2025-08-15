"""
Universal File Parser System for ESG Data Extraction
Converts any uploaded file (PDF, Excel, CSV, Images, etc.) to structured JSON
"""

import json
import re
import logging
from pathlib import Path
from typing import Dict, Any, List, Optional, Union
from datetime import datetime
import mimetypes

# Document parsing libraries - with fallbacks for missing dependencies
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    from PIL import Image
    import pytesseract
except ImportError:
    Image = None
    pytesseract = None

try:
    import docx
except ImportError:
    docx = None

import csv
import xml.etree.ElementTree as ET

# For advanced text extraction and NLP
from dataclasses import dataclass, asdict
import dateutil.parser

logger = logging.getLogger(__name__)


@dataclass
class ExtractedData:
    """Standardized structure for extracted data"""
    # Metadata
    file_name: str
    file_type: str
    extraction_date: str
    extraction_method: str
    confidence_score: float
    
    # Environmental metrics
    energy_consumption_kwh: Optional[float] = None
    water_usage_liters: Optional[float] = None
    waste_generated_kg: Optional[float] = None
    carbon_emissions_tco2: Optional[float] = None
    renewable_energy_percentage: Optional[float] = None
    
    # Social metrics
    total_employees: Optional[int] = None
    training_hours: Optional[float] = None
    safety_incidents: Optional[int] = None
    employee_satisfaction_score: Optional[float] = None
    diversity_ratio: Optional[float] = None
    
    # Governance metrics
    board_meetings: Optional[int] = None
    compliance_score: Optional[float] = None
    audit_findings: Optional[int] = None
    policy_updates: Optional[int] = None
    
    # Additional extracted data
    raw_text: Optional[str] = None
    tables: Optional[List[Dict]] = None
    key_value_pairs: Optional[Dict[str, Any]] = None
    dates_found: Optional[List[str]] = None
    amounts_found: Optional[List[Dict[str, float]]] = None
    
    def to_json(self) -> str:
        """Convert to JSON string"""
        return json.dumps(asdict(self), default=str, indent=2)


class UniversalFileParser:
    """
    Universal file parser that handles multiple file types
    and extracts ESG-relevant data
    """
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._parse_pdf,
            '.xlsx': self._parse_excel,
            '.xls': self._parse_excel,
            '.csv': self._parse_csv,
            '.docx': self._parse_docx,
            '.doc': self._parse_doc,
            '.png': self._parse_image,
            '.jpg': self._parse_image,
            '.jpeg': self._parse_image,
            '.txt': self._parse_text,
            '.json': self._parse_json,
            '.xml': self._parse_xml,
        }
        
        # ESG keyword patterns for intelligent extraction
        self.esg_patterns = {
            'energy': {
                'keywords': ['kwh', 'kilowatt', 'energy', 'electricity', 'power consumption'],
                'unit_multipliers': {'kwh': 1, 'mwh': 1000, 'gwh': 1000000},
                'pattern': r'(\d+(?:,\d{3})*(?:\.\d+)?)\s*(?:kwh|mwh|kilowatt)',
            },
            'water': {
                'keywords': ['water', 'liters', 'gallons', 'cubic meters', 'm³'],
                'unit_multipliers': {'liters': 1, 'gallons': 3.785, 'm3': 1000, 'cubic meters': 1000},
                'pattern': r'(\d+(?:,\d{3})*(?:\.\d+)?)\s*(?:liters?|gallons?|m³|cubic meters?)',
            },
            'waste': {
                'keywords': ['waste', 'kg', 'tons', 'disposal', 'recycling'],
                'unit_multipliers': {'kg': 1, 'ton': 1000, 'metric ton': 1000},
                'pattern': r'(\d+(?:,\d{3})*(?:\.\d+)?)\s*(?:kg|tons?|metric tons?)',
            },
            'carbon': {
                'keywords': ['co2', 'carbon', 'emissions', 'tco2', 'greenhouse'],
                'unit_multipliers': {'tco2': 1, 'kg co2': 0.001, 'metric tons': 1},
                'pattern': r'(\d+(?:,\d{3})*(?:\.\d+)?)\s*(?:tco2|tons? co2|metric tons?)',
            },
            'employees': {
                'keywords': ['employees', 'staff', 'workforce', 'headcount'],
                'pattern': r'(?:total\s+)?(?:employees?|staff|workforce)[:\s]+(\d+)',
            },
            'training': {
                'keywords': ['training', 'hours', 'development', 'learning'],
                'pattern': r'training\s+hours?[:\s]+(\d+(?:,\d{3})*(?:\.\d+)?)',
            },
        }
    
    def parse_file(self, file_path: str, file_obj=None) -> ExtractedData:
        """
        Main entry point for parsing any file
        
        Args:
            file_path: Path to the file or filename
            file_obj: File object (for Django FileField)
        
        Returns:
            ExtractedData object with all extracted information
        """
        try:
            # Determine file type
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension not in self.supported_formats:
                logger.warning(f"Unsupported file format: {file_extension}")
                return self._create_empty_result(file_path, "unsupported")
            
            # Call appropriate parser
            parser_func = self.supported_formats[file_extension]
            
            if file_obj:
                # For Django FileField objects
                return parser_func(file_obj, file_path)
            else:
                # For file paths
                with open(file_path, 'rb') as f:
                    return parser_func(f, file_path)
                    
        except Exception as e:
            logger.error(f"Error parsing file {file_path}: {e}")
            return self._create_empty_result(file_path, "error", str(e))
    
    def _parse_pdf(self, file_obj, file_name: str) -> ExtractedData:
        """Parse PDF files using multiple methods for best results"""
        extracted_data = ExtractedData(
            file_name=file_name,
            file_type='pdf',
            extraction_date=datetime.now().isoformat(),
            extraction_method='pdf_extraction',
            confidence_score=0.0
        )
        
        if not pdfplumber and not PyPDF2:
            extracted_data.raw_text = "PDF parsing libraries not available"
            return extracted_data
        
        try:
            # Method 1: Try pdfplumber first (better for tables)
            text = ""
            tables = []
            
            if pdfplumber:
                with pdfplumber.open(file_obj) as pdf:
                    for page in pdf.pages:
                        # Extract text
                        page_text = page.extract_text() or ""
                        text += page_text + "\n"
                        
                        # Extract tables
                        page_tables = page.extract_tables()
                        if page_tables:
                            for table in page_tables:
                                if table:
                                    # Convert table to dict format
                                    if len(table) > 1:
                                        headers = table[0]
                                        for row in table[1:]:
                                            if row and headers:
                                                table_dict = dict(zip(headers, row))
                                                tables.append(table_dict)
            
            # Method 2: Fallback to PyPDF2 if needed
            if not text.strip() and PyPDF2:
                file_obj.seek(0)
                pdf_reader = PyPDF2.PdfReader(file_obj)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            
            extracted_data.raw_text = text
            extracted_data.tables = tables
            
            # Extract ESG metrics from text
            self._extract_esg_metrics(extracted_data, text)
            
            # Extract from tables
            if tables:
                self._extract_from_tables(extracted_data, tables)
            
            # Calculate confidence score
            extracted_data.confidence_score = self._calculate_confidence(extracted_data)
            
        except Exception as e:
            logger.error(f"Error parsing PDF {file_name}: {e}")
            # Try to parse as text if PDF parsing fails
            try:
                file_obj.seek(0)
                content = file_obj.read()
                if isinstance(content, bytes):
                    text = content.decode('utf-8', errors='ignore')
                else:
                    text = str(content)
                
                extracted_data.raw_text = text
                extracted_data.extraction_method = 'text_fallback'
                
                # Extract ESG metrics from the text
                self._extract_esg_metrics(extracted_data, text)
                
                # Recalculate confidence
                extracted_data.confidence_score = self._calculate_confidence(extracted_data)
                
                logger.info(f"Successfully parsed {file_name} as text after PDF parsing failed")
            except Exception as text_error:
                logger.error(f"Text fallback also failed for {file_name}: {text_error}")
                extracted_data.raw_text = f"Error: {str(e)}"
        
        return extracted_data
    
    def _parse_excel(self, file_obj, file_name: str) -> ExtractedData:
        """Parse Excel files and extract data from all sheets"""
        extracted_data = ExtractedData(
            file_name=file_name,
            file_type='excel',
            extraction_date=datetime.now().isoformat(),
            extraction_method='excel_extraction',
            confidence_score=0.0
        )
        
        if not pd:
            extracted_data.raw_text = "Excel parsing libraries not available"
            return extracted_data
        
        try:
            # Read Excel file
            excel_data = pd.read_excel(file_obj, sheet_name=None)
            
            all_tables = []
            all_text = []
            
            for sheet_name, df in excel_data.items():
                # Convert DataFrame to dict records
                sheet_data = df.to_dict('records')
                all_tables.extend(sheet_data)
                
                # Also extract as text for pattern matching
                sheet_text = df.to_string()
                all_text.append(f"Sheet: {sheet_name}\n{sheet_text}")
                
                # Look for specific ESG metrics in known columns
                self._extract_from_dataframe(extracted_data, df)
            
            extracted_data.tables = all_tables
            extracted_data.raw_text = "\n\n".join(all_text)
            
            # Extract metrics from combined text
            self._extract_esg_metrics(extracted_data, extracted_data.raw_text)
            
            extracted_data.confidence_score = self._calculate_confidence(extracted_data)
            
        except Exception as e:
            logger.error(f"Error parsing Excel {file_name}: {e}")
            extracted_data.raw_text = f"Error: {str(e)}"
        
        return extracted_data
    
    def _parse_csv(self, file_obj, file_name: str) -> ExtractedData:
        """Parse CSV files"""
        extracted_data = ExtractedData(
            file_name=file_name,
            file_type='csv',
            extraction_date=datetime.now().isoformat(),
            extraction_method='csv_extraction',
            confidence_score=0.0
        )
        
        if not pd:
            extracted_data.raw_text = "CSV parsing libraries not available"
            return extracted_data
        
        try:
            # Try to detect encoding
            file_obj.seek(0)
            content = file_obj.read()
            
            # Try UTF-8 first, then fallback to latin-1
            for encoding in ['utf-8', 'latin-1', 'iso-8859-1']:
                try:
                    text = content.decode(encoding)
                    break
                except:
                    continue
            
            # Parse CSV
            df = pd.read_csv(file_obj)
            
            # Convert to records
            extracted_data.tables = df.to_dict('records')
            extracted_data.raw_text = df.to_string()
            
            # Extract from DataFrame
            self._extract_from_dataframe(extracted_data, df)
            
            # Extract from text
            self._extract_esg_metrics(extracted_data, extracted_data.raw_text)
            
            extracted_data.confidence_score = self._calculate_confidence(extracted_data)
            
        except Exception as e:
            logger.error(f"Error parsing CSV {file_name}: {e}")
            extracted_data.raw_text = f"Error: {str(e)}"
        
        return extracted_data
    
    def _parse_image(self, file_obj, file_name: str) -> ExtractedData:
        """Parse images using OCR"""
        extracted_data = ExtractedData(
            file_name=file_name,
            file_type='image',
            extraction_date=datetime.now().isoformat(),
            extraction_method='ocr_extraction',
            confidence_score=0.0
        )
        
        if not Image or not pytesseract:
            extracted_data.raw_text = "Image OCR libraries not available"
            return extracted_data
        
        try:
            # Open image
            image = Image.open(file_obj)
            
            # Perform OCR
            text = pytesseract.image_to_string(image)
            extracted_data.raw_text = text
            
            # Try to extract data with OCR confidence
            ocr_data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
            
            # Extract ESG metrics from OCR text
            self._extract_esg_metrics(extracted_data, text)
            
            # Lower confidence for OCR
            extracted_data.confidence_score = self._calculate_confidence(extracted_data) * 0.7
            
        except Exception as e:
            logger.error(f"Error parsing image {file_name}: {e}")
            extracted_data.raw_text = f"Error: {str(e)}"
        
        return extracted_data
    
    def _parse_docx(self, file_obj, file_name: str) -> ExtractedData:
        """Parse Word documents"""
        extracted_data = ExtractedData(
            file_name=file_name,
            file_type='docx',
            extraction_date=datetime.now().isoformat(),
            extraction_method='docx_extraction',
            confidence_score=0.0
        )
        
        if not docx:
            extracted_data.raw_text = "DOCX parsing libraries not available"
            return extracted_data
        
        try:
            doc = docx.Document(file_obj)
            
            # Extract text from paragraphs
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            
            # Extract text from tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    # Convert to dict format
                    if len(table_data) > 1:
                        headers = table_data[0]
                        for row in table_data[1:]:
                            table_dict = dict(zip(headers, row))
                            tables.append(table_dict)
            
            extracted_data.raw_text = "\n".join(text)
            extracted_data.tables = tables
            
            # Extract ESG metrics
            self._extract_esg_metrics(extracted_data, extracted_data.raw_text)
            
            # Extract from tables
            if tables:
                self._extract_from_tables(extracted_data, tables)
            
            extracted_data.confidence_score = self._calculate_confidence(extracted_data)
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_name}: {e}")
            extracted_data.raw_text = f"Error: {str(e)}"
        
        return extracted_data
    
    def _parse_doc(self, file_obj, file_name: str) -> ExtractedData:
        """Parse old Word documents (fallback to text extraction)"""
        # For .doc files, we might need python-docx2txt or similar
        # For now, treat as binary and try text extraction
        return self._parse_text(file_obj, file_name)
    
    def _parse_text(self, file_obj, file_name: str) -> ExtractedData:
        """Parse plain text files"""
        extracted_data = ExtractedData(
            file_name=file_name,
            file_type='text',
            extraction_date=datetime.now().isoformat(),
            extraction_method='text_extraction',
            confidence_score=0.0
        )
        
        try:
            # Read text
            if hasattr(file_obj, 'read'):
                content = file_obj.read()
                if isinstance(content, bytes):
                    text = content.decode('utf-8', errors='ignore')
                else:
                    text = content
            else:
                text = str(file_obj)
            
            extracted_data.raw_text = text
            
            # Extract ESG metrics
            self._extract_esg_metrics(extracted_data, text)
            
            extracted_data.confidence_score = self._calculate_confidence(extracted_data)
            
        except Exception as e:
            logger.error(f"Error parsing text {file_name}: {e}")
            extracted_data.raw_text = f"Error: {str(e)}"
        
        return extracted_data
    
    def _parse_json(self, file_obj, file_name: str) -> ExtractedData:
        """Parse JSON files"""
        extracted_data = ExtractedData(
            file_name=file_name,
            file_type='json',
            extraction_date=datetime.now().isoformat(),
            extraction_method='json_extraction',
            confidence_score=0.0
        )
        
        try:
            # Load JSON
            if hasattr(file_obj, 'read'):
                json_data = json.load(file_obj)
            else:
                json_data = json.loads(file_obj)
            
            # Store as key-value pairs
            extracted_data.key_value_pairs = self._flatten_json(json_data)
            extracted_data.raw_text = json.dumps(json_data, indent=2)
            
            # Try to map JSON fields to ESG metrics
            self._extract_from_json(extracted_data, json_data)
            
            extracted_data.confidence_score = self._calculate_confidence(extracted_data)
            
        except Exception as e:
            logger.error(f"Error parsing JSON {file_name}: {e}")
            extracted_data.raw_text = f"Error: {str(e)}"
        
        return extracted_data
    
    def _parse_xml(self, file_obj, file_name: str) -> ExtractedData:
        """Parse XML files"""
        extracted_data = ExtractedData(
            file_name=file_name,
            file_type='xml',
            extraction_date=datetime.now().isoformat(),
            extraction_method='xml_extraction',
            confidence_score=0.0
        )
        
        try:
            # Parse XML
            tree = ET.parse(file_obj)
            root = tree.getroot()
            
            # Convert to dict
            xml_dict = self._xml_to_dict(root)
            
            extracted_data.key_value_pairs = self._flatten_json(xml_dict)
            extracted_data.raw_text = ET.tostring(root, encoding='unicode')
            
            # Extract from flattened data
            self._extract_from_json(extracted_data, xml_dict)
            
            extracted_data.confidence_score = self._calculate_confidence(extracted_data)
            
        except Exception as e:
            logger.error(f"Error parsing XML {file_name}: {e}")
            extracted_data.raw_text = f"Error: {str(e)}"
        
        return extracted_data
    
    def _extract_esg_metrics(self, data: ExtractedData, text: str) -> None:
        """Extract ESG metrics from text using patterns"""
        if not text:
            return
        
        text_lower = text.lower()
        
        # Extract energy consumption
        for match in re.finditer(self.esg_patterns['energy']['pattern'], text_lower):
            value = self._parse_number(match.group(1))
            if value and not data.energy_consumption_kwh:
                data.energy_consumption_kwh = value
        
        # Extract water usage
        for match in re.finditer(self.esg_patterns['water']['pattern'], text_lower):
            value = self._parse_number(match.group(1))
            if value and not data.water_usage_liters:
                data.water_usage_liters = value
        
        # Extract waste
        for match in re.finditer(self.esg_patterns['waste']['pattern'], text_lower):
            value = self._parse_number(match.group(1))
            if value and not data.waste_generated_kg:
                data.waste_generated_kg = value
        
        # Extract carbon emissions
        for match in re.finditer(self.esg_patterns['carbon']['pattern'], text_lower):
            value = self._parse_number(match.group(1))
            if value and not data.carbon_emissions_tco2:
                data.carbon_emissions_tco2 = value
        
        # Extract employee count
        for match in re.finditer(self.esg_patterns['employees']['pattern'], text_lower):
            value = self._parse_number(match.group(1))
            if value and not data.total_employees:
                data.total_employees = int(value)
        
        # Extract training hours
        for match in re.finditer(self.esg_patterns['training']['pattern'], text_lower):
            value = self._parse_number(match.group(1))
            if value and not data.training_hours:
                data.training_hours = value
        
        # Extract percentages (for scores and ratios)
        percentage_pattern = r'(\d+(?:\.\d+)?)\s*%'
        percentages = []
        for match in re.finditer(percentage_pattern, text):
            value = float(match.group(1))
            percentages.append(value)
        
        # Try to assign percentages to appropriate fields
        if percentages:
            # Look for context around percentages
            if 'renewable' in text_lower and not data.renewable_energy_percentage:
                for pct in percentages:
                    if 0 <= pct <= 100:
                        data.renewable_energy_percentage = pct
                        break
            
            if 'satisfaction' in text_lower and not data.employee_satisfaction_score:
                for pct in percentages:
                    if 50 <= pct <= 100:
                        data.employee_satisfaction_score = pct
                        break
            
            if 'compliance' in text_lower and not data.compliance_score:
                for pct in percentages:
                    if 0 <= pct <= 100:
                        data.compliance_score = pct
                        break
        
        # Extract dates
        date_pattern = r'\b(\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}[/-]\d{1,2}[/-]\d{1,2})\b'
        dates = re.findall(date_pattern, text)
        if dates:
            data.dates_found = dates[:10]  # Limit to 10 dates
        
        # Extract monetary amounts
        money_pattern = r'(?:AED|USD|EUR|GBP)?\s*(\d+(?:,\d{3})*(?:\.\d{2})?)'
        amounts = []
        for match in re.finditer(money_pattern, text):
            value = self._parse_number(match.group(1))
            if value:
                amounts.append({'value': value, 'context': text[max(0, match.start()-20):match.end()+20]})
        
        if amounts:
            data.amounts_found = amounts[:10]  # Limit to 10 amounts
    
    def _extract_from_tables(self, data: ExtractedData, tables: List[Dict]) -> None:
        """Extract ESG metrics from table data"""
        for table in tables:
            for key, value in table.items():
                if not key or not value:
                    continue
                
                key_lower = str(key).lower()
                value_str = str(value).lower()
                
                # Try to match with ESG fields
                if 'energy' in key_lower or 'kwh' in key_lower:
                    parsed_value = self._parse_number(value)
                    if parsed_value and not data.energy_consumption_kwh:
                        data.energy_consumption_kwh = parsed_value
                
                elif 'water' in key_lower:
                    parsed_value = self._parse_number(value)
                    if parsed_value and not data.water_usage_liters:
                        data.water_usage_liters = parsed_value
                
                elif 'waste' in key_lower:
                    parsed_value = self._parse_number(value)
                    if parsed_value and not data.waste_generated_kg:
                        data.waste_generated_kg = parsed_value
                
                elif 'carbon' in key_lower or 'co2' in key_lower or 'emission' in key_lower:
                    parsed_value = self._parse_number(value)
                    if parsed_value and not data.carbon_emissions_tco2:
                        data.carbon_emissions_tco2 = parsed_value
                
                elif 'employee' in key_lower and ('count' in key_lower or 'total' in key_lower):
                    parsed_value = self._parse_number(value)
                    if parsed_value and not data.total_employees:
                        data.total_employees = int(parsed_value)
                
                elif 'training' in key_lower and 'hour' in key_lower:
                    parsed_value = self._parse_number(value)
                    if parsed_value and not data.training_hours:
                        data.training_hours = parsed_value
    
    def _extract_from_dataframe(self, data: ExtractedData, df: pd.DataFrame) -> None:
        """Extract ESG metrics from pandas DataFrame"""
        # Look for columns with ESG-related names
        for col in df.columns:
            col_lower = str(col).lower()
            
            # Skip if column has no data
            if df[col].isna().all():
                continue
            
            # Try to extract based on column names
            if 'energy' in col_lower or 'kwh' in col_lower:
                try:
                    value = pd.to_numeric(df[col], errors='coerce').sum()
                    if value > 0 and not data.energy_consumption_kwh:
                        data.energy_consumption_kwh = float(value)
                except:
                    pass
            
            elif 'water' in col_lower:
                try:
                    value = pd.to_numeric(df[col], errors='coerce').sum()
                    if value > 0 and not data.water_usage_liters:
                        data.water_usage_liters = float(value)
                except:
                    pass
            
            elif 'waste' in col_lower:
                try:
                    value = pd.to_numeric(df[col], errors='coerce').sum()
                    if value > 0 and not data.waste_generated_kg:
                        data.waste_generated_kg = float(value)
                except:
                    pass
            
            elif 'employee' in col_lower:
                try:
                    # For employee count, take the max or last value
                    value = pd.to_numeric(df[col], errors='coerce').max()
                    if value > 0 and not data.total_employees:
                        data.total_employees = int(value)
                except:
                    pass
    
    def _extract_from_json(self, data: ExtractedData, json_data: Dict) -> None:
        """Extract ESG metrics from JSON data"""
        # Flatten JSON and look for ESG fields
        flat_data = self._flatten_json(json_data)
        
        for key, value in flat_data.items():
            key_lower = str(key).lower()
            
            # Map to ESG fields
            if 'energy' in key_lower or 'kwh' in key_lower:
                parsed_value = self._parse_number(value)
                if parsed_value and not data.energy_consumption_kwh:
                    data.energy_consumption_kwh = parsed_value
            
            elif 'water' in key_lower:
                parsed_value = self._parse_number(value)
                if parsed_value and not data.water_usage_liters:
                    data.water_usage_liters = parsed_value
            
            elif 'waste' in key_lower:
                parsed_value = self._parse_number(value)
                if parsed_value and not data.waste_generated_kg:
                    data.waste_generated_kg = parsed_value
            
            elif 'carbon' in key_lower or 'co2' in key_lower:
                parsed_value = self._parse_number(value)
                if parsed_value and not data.carbon_emissions_tco2:
                    data.carbon_emissions_tco2 = parsed_value
            
            elif 'employee' in key_lower:
                parsed_value = self._parse_number(value)
                if parsed_value and not data.total_employees:
                    data.total_employees = int(parsed_value)
    
    def _parse_number(self, value: Any) -> Optional[float]:
        """Parse number from various formats"""
        if value is None:
            return None
        
        if isinstance(value, (int, float)):
            return float(value)
        
        # Convert to string and clean
        value_str = str(value)
        
        # Remove common separators and currency symbols
        value_str = value_str.replace(',', '')
        value_str = value_str.replace('AED', '')
        value_str = value_str.replace('USD', '')
        value_str = value_str.replace('$', '')
        value_str = value_str.strip()
        
        try:
            return float(value_str)
        except:
            # Try to extract first number
            match = re.search(r'(\d+(?:\.\d+)?)', value_str)
            if match:
                try:
                    return float(match.group(1))
                except:
                    pass
        
        return None
    
    def _flatten_json(self, json_data: Any, parent_key: str = '', sep: str = '.') -> Dict:
        """Flatten nested JSON structure"""
        items = []
        
        if isinstance(json_data, dict):
            for k, v in json_data.items():
                new_key = f"{parent_key}{sep}{k}" if parent_key else k
                if isinstance(v, (dict, list)):
                    items.extend(self._flatten_json(v, new_key, sep=sep).items())
                else:
                    items.append((new_key, v))
        elif isinstance(json_data, list):
            for i, v in enumerate(json_data):
                new_key = f"{parent_key}[{i}]"
                if isinstance(v, (dict, list)):
                    items.extend(self._flatten_json(v, new_key, sep=sep).items())
                else:
                    items.append((new_key, v))
        else:
            items.append((parent_key, json_data))
        
        return dict(items)
    
    def _xml_to_dict(self, element) -> Dict:
        """Convert XML element to dictionary"""
        result = {}
        
        # Add attributes
        if element.attrib:
            result['@attributes'] = element.attrib
        
        # Add text content
        if element.text and element.text.strip():
            result['text'] = element.text.strip()
        
        # Add children
        children = {}
        for child in element:
            child_data = self._xml_to_dict(child)
            if child.tag in children:
                # Multiple children with same tag
                if not isinstance(children[child.tag], list):
                    children[child.tag] = [children[child.tag]]
                children[child.tag].append(child_data)
            else:
                children[child.tag] = child_data
        
        if children:
            result.update(children)
        
        return result if result else element.text
    
    def _calculate_confidence(self, data: ExtractedData) -> float:
        """Calculate confidence score based on extracted data"""
        score = 0.0
        fields_checked = 0
        
        # Check environmental metrics
        if data.energy_consumption_kwh is not None:
            score += 1
        fields_checked += 1
        
        if data.water_usage_liters is not None:
            score += 1
        fields_checked += 1
        
        if data.waste_generated_kg is not None:
            score += 1
        fields_checked += 1
        
        if data.carbon_emissions_tco2 is not None:
            score += 1
        fields_checked += 1
        
        # Check social metrics
        if data.total_employees is not None:
            score += 1
        fields_checked += 1
        
        if data.training_hours is not None:
            score += 1
        fields_checked += 1
        
        # Check governance metrics
        if data.compliance_score is not None:
            score += 1
        fields_checked += 1
        
        # Check if we have raw text
        if data.raw_text and len(data.raw_text) > 100:
            score += 0.5
        
        # Check if we have tables
        if data.tables and len(data.tables) > 0:
            score += 0.5
        
        # Calculate percentage
        if fields_checked > 0:
            confidence = (score / (fields_checked + 1)) * 100
        else:
            confidence = 0.0
        
        return min(confidence, 100.0)
    
    def _create_empty_result(self, file_name: str, status: str, error: str = None) -> ExtractedData:
        """Create empty result for unsupported or errored files"""
        return ExtractedData(
            file_name=file_name,
            file_type=status,
            extraction_date=datetime.now().isoformat(),
            extraction_method='none',
            confidence_score=0.0,
            raw_text=error or f"File type {status}"
        )